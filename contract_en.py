import streamlit as st
from docx import Document
import io
import datetime

def replace_placeholder(paragraph, key, value):
    """Improved placeholder replacement preserving formatting"""
    if key not in paragraph.text:
        return
    
    # Store all runs and their text
    runs = list(paragraph.runs)
    text = ''.join(run.text for run in runs)
    
    if key not in text:
        return
    
    # Find start and end positions of the key
    start_idx = text.find(key)
    end_idx = start_idx + len(key)
    
    # Clear paragraph and rebuild
    paragraph.clear()
    current_idx = 0
    
    for run in runs:
        run_text = run.text
        run_end = current_idx + len(run_text)
        
        # Case 1: Run is entirely before the key
        if run_end <= start_idx:
            new_run = paragraph.add_run(run_text)
            copy_run_formatting(new_run, run)
        
        # Case 2: Run contains start of key
        elif current_idx <= start_idx < run_end:
            # Part before key
            if current_idx < start_idx:
                prefix = run_text[:start_idx - current_idx]
                new_run = paragraph.add_run(prefix)
                copy_run_formatting(new_run, run)
            
            # Replacement text with original formatting
            new_run = paragraph.add_run(value)
            copy_run_formatting(new_run, run)
            
            # Part after key (if any)
            post_start = start_idx - current_idx + len(key)
            if post_start < len(run_text):
                suffix = run_text[post_start:]
                new_run = paragraph.add_run(suffix)
                copy_run_formatting(new_run, run)
        
        # Case 3: Run is within key (skip)
        elif start_idx < current_idx < end_idx:
            pass
        
        # Case 4: Run after key
        else:
            suffix_start = max(0, end_idx - current_idx)
            suffix = run_text[suffix_start:]
            if suffix:
                new_run = paragraph.add_run(suffix)
                copy_run_formatting(new_run, run)
        
        current_idx = run_end

def copy_run_formatting(new_run, original_run):
    """Copy formatting from original run to new run, safely handling missing properties"""
    # Basic font styles
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic
    new_run.underline = original_run.underline
    new_run.font.size = original_run.font.size

    # Font name (Latin)
    if original_run.font.name:
        new_run.font.name = original_run.font.name

    # Copy East Asian fonts if present
    rPr = getattr(original_run._element, "rPr", None)
    if rPr is not None:
        rFonts = getattr(rPr, "rFonts", None)
        if rFonts is not None and getattr(rFonts, "eastAsia", None):
            new_run._element.rPr.rFonts = rFonts

# --- Streamlit UI ---
st.title("📝 Contract Generator")
st.write(f"Created by **Dimitris Katsanos**")

with st.form("contract_form"):
    name = st.text_input("Full Name")
    company = st.text_input("Company Name")
    email = st.text_input("E-mail")
    telephone = st.text_input("Telephone")
    date = st.date_input("Agreement Date", datetime.date.today())
    terms = st.text_area("Data Usage", "Data will be used for the elaboration...")
    item1= st.text_input("item number e.g. 1")
    data1= st.text_input("Parameter e.g. Temperature")
    peri1= st.text_input("Time Step & Duration e.g. Daily Values, 1981 - 2000")
    item2= st.text_input("item number e.g. 2")
    data2= st.text_input("Parameter e.g. Humidity")
    peri2= st.text_input("Time Step & Duration e.g. Hourly Values, 1981 - 2000")
    item3= st.text_input("item number e.g. 3")
    data3=st.text_input("Parameter e.g. Pressure")
    peri3= st.text_input("Time Step & Duration e.g. Monthly Values, 1981 - 2000")
    name2 = st.text_input("Supervisor (if applicable)")
    submit_button = st.form_submit_button("Generate Contract")

if submit_button:
    doc = Document("template.docx")
    formatted_date = date.strftime("%d/%m/%Y")
    
    replacements = {
        "{{name}}": name,
        "{{company}}": company,
        "{{email}}": email,
        "{{telephone}}": telephone,
        "{{date}}": formatted_date,
        "{{data1}}": data1,
        "{{data2}}": data2,
        "{{data3}}": data3,
        "{{item1}}": item1,
        "{{item2}}": item2,
        "{{item3}}": item3,
        "{{peri1}}": peri1,
        "{{peri2}}": peri2,
        "{{peri3}}": peri3,
        "{{name2}}": name2,
        "{{terms}}": terms
    }
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            replace_placeholder(paragraph, key, value)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_placeholder(paragraph, key, value)
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    st.success("Contract generated!")
    st.download_button(
        label="📥 Download Contract",
        data=doc_bytes,
        file_name=f"Contract_{name.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
