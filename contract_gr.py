import streamlit as st
from docx import Document
import io
import datetime

def replace_placeholder(paragraph, key, value):
    """Improved placeholder replacement preserving formatting"""
    if key not in paragraph.text:
        return
    
    # Store all runs and their text
    runs = list(paragraph.runs)
    text = ''.join(run.text for run in runs)
    
    if key not in text:
        return
    
    # Find start and end positions of the key
    start_idx = text.find(key)
    end_idx = start_idx + len(key)
    
    # Clear paragraph and rebuild
    paragraph.clear()
    current_idx = 0
    
    for run in runs:
        run_text = run.text
        run_end = current_idx + len(run_text)
        
        # Case 1: Run is entirely before the key
        if run_end <= start_idx:
            new_run = paragraph.add_run(run_text)
            copy_run_formatting(new_run, run)
        
        # Case 2: Run contains start of key
        elif current_idx <= start_idx < run_end:
            # Part before key
            if current_idx < start_idx:
                prefix = run_text[:start_idx - current_idx]
                new_run = paragraph.add_run(prefix)
                copy_run_formatting(new_run, run)
            
            # Replacement text with original formatting
            new_run = paragraph.add_run(value)
            copy_run_formatting(new_run, run)
            
            # Part after key (if any)
            post_start = start_idx - current_idx + len(key)
            if post_start < len(run_text):
                suffix = run_text[post_start:]
                new_run = paragraph.add_run(suffix)
                copy_run_formatting(new_run, run)
        
        # Case 3: Run is within key (skip)
        elif start_idx < current_idx < end_idx:
            pass
        
        # Case 4: Run after key
        else:
            suffix_start = max(0, end_idx - current_idx)
            suffix = run_text[suffix_start:]
            if suffix:
                new_run = paragraph.add_run(suffix)
                copy_run_formatting(new_run, run)
        
        current_idx = run_end

def copy_run_formatting(new_run, original_run):
    """Copy formatting from original run to new run, safely handling missing properties"""
    # Basic font styles
    new_run.bold = original_run.bold
    new_run.italic = original_run.italic
    new_run.underline = original_run.underline
    new_run.font.size = original_run.font.size

    # Font name (Latin)
    if original_run.font.name:
        new_run.font.name = original_run.font.name

    # Copy East Asian fonts if present
    rPr = getattr(original_run._element, "rPr", None)
    if rPr is not None:
        rFonts = getattr(rPr, "rFonts", None)
        if rFonts is not None and getattr(rFonts, "eastAsia", None):
            new_run._element.rPr.rFonts = rFonts

# --- Streamlit UI ---
st.title("ðŸ“ Î”Î·Î¼Î¹Î¿Ï…ÏÎ³Î¯Î± Î£Ï…Î¼Ï†Ï‰Î½Î·Ï„Î¹ÎºÎ¿Ï")
st.write(f"Created by **Dimitris Katsanos**")

with st.form("contract_form"):
    name = st.text_input("ÎŸÎ½Î¿Î¼Î± & Î•Ï€ÏŽÎ½Ï…Î¼Î¿")
    company = st.text_input("ÎŸÏÎ³Î±Î½Î¹ÏƒÎ¼ÏŒÏ‚/Î•Ï„Î±Î¹ÏÎ¯Î±")
    email = st.text_input("E-mail")
    telephone = st.text_input("Î¤Î·Î»Î­Ï†Ï‰Î½Î¿")
    date = st.date_input("Î—Î¼ÎµÏÎ¿Î¼Î·Î½Î¯Î±", datetime.date.today())
    terms = st.text_area("Î§ÏÎ®ÏƒÎ· Î”ÎµÎ´Î¿Î¼Î­Î½Ï‰Î½", "Î¤Î± Î´ÎµÎ´Î¿Î¼Î­Î½Î± Î¸Î± Ï‡ÏÎ·ÏƒÎ¼Î¿Ï€Î¿Î¹Î·Î¸Î¿ÏÎ½ ...")
    item1= st.text_input("Î±/Î± ÏƒÏ„Î¿Î¹Ï‡ÎµÎ¯Î¿Ï… Ï€.Ï‡. 1")
    data1= st.text_input("Î Î±ÏÎ¬Î¼ÎµÏ„ÏÎ¿Ï‚ Ï€.Ï‡. Î˜ÎµÏÎ¼Î¿ÎºÏÎ±ÏƒÎ¯Î±")
    peri1= st.text_input("Î§ÏÎ¿Î½Î¹ÎºÎ® Î‘Î½Î¬Î»Ï…ÏƒÎ· & Î”Î¹Î¬ÏÎºÎµÎ¹Î± Ï€.Ï‡. Î—Î¼ÎµÏÎ®ÏƒÎ¹ÎµÏ‚ Ï„Î¹Î¼Î­Ï‚, 1981 - 2000")
    item2= st.text_input("Î±/Î± ÏƒÏ„Î¿Î¹Ï‡ÎµÎ¯Î¿Ï… Ï€.Ï‡. 2")
    data2= st.text_input("Î Î±ÏÎ¬Î¼ÎµÏ„ÏÎ¿Ï‚ Ï€.Ï‡. Î¥Î³ÏÎ±ÏƒÎ¯Î±")
    peri2= st.text_input("Î§ÏÎ¿Î½Î¹ÎºÎ® Î‘Î½Î¬Î»Ï…ÏƒÎ· & Î”Î¹Î¬ÏÎºÎµÎ¹Î± Ï€.Ï‡. Î©ÏÎ¹Î±Î¯ÎµÏ‚ Ï„Î¹Î¼Î­Ï‚, 1981 - 2000")
    item3= st.text_input("Î±/Î± ÏƒÏ„Î¿Î¹Ï‡ÎµÎ¯Î¿Ï… Ï€.Ï‡. 3")
    data3=st.text_input("Î Î±ÏÎ¬Î¼ÎµÏ„ÏÎ¿Ï‚ Ï€.Ï‡. Î Î¯ÎµÏƒÎ·")
    peri3= st.text_input("Î§ÏÎ¿Î½Î¹ÎºÎ® Î‘Î½Î¬Î»Ï…ÏƒÎ· & Î”Î¹Î¬ÏÎºÎµÎ¹Î± Ï€.Ï‡. ÎœÎ·Î½Î¹Î±Î¯ÎµÏ‚ Ï„Î¹Î¼Î­Ï‚, 1981 - 2000")
    name2 = st.text_input("Î•Ï€Î¹Î²Î»Î­Ï€Ï‰Î½ ÎšÎ±Î¸Î·Î³Î·Ï„Î®Ï‚ (Î±Î½ Ï…Ï€Î¬ÏÏ‡ÎµÎ¹)")
    submit_button = st.form_submit_button("Î”Î·Î¼Î¹Î¿Ï…ÏÎ³Î¯Î± Î£Ï…Î¼Ï†Ï‰Î½Î·Ï„Î¹ÎºÎ¿Ï")

if submit_button:
    doc = Document("template2.docx")
    formatted_date = date.strftime("%d/%m/%Y")
    
    replacements = {
        "{{name}}": name,
        "{{company}}": company,
        "{{email}}": email,
        "{{telephone}}": telephone,
        "{{date}}": formatted_date,
        "{{data1}}": data1,
        "{{data2}}": data2,
        "{{data3}}": data3,
        "{{item1}}": item1,
        "{{item2}}": item2,
        "{{item3}}": item3,
        "{{peri1}}": peri1,
        "{{peri2}}": peri2,
        "{{peri3}}": peri3,
        "{{name2}}": name2,
        "{{terms}}": terms
    }
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            replace_placeholder(paragraph, key, value)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_placeholder(paragraph, key, value)
    
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    st.success("Contract generated!")
    st.download_button(
        label="ðŸ“¥ Download Contract",
        data=doc_bytes,
        file_name=f"Contract_{name.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )